import docx
from fpdf import FPDF
from app.models.schemas import QuizQuestion
from typing import List

def export_docx(quiz_data: List[QuizQuestion], filepath: str, role: str = "Teacher"):
    doc = docx.Document()
    doc.add_heading('Generated Challenge Quiz', 0)
    
    # Write all questions first (without answers/explanations)
    for i, q in enumerate(quiz_data, 1):
        doc.add_heading(f"Q{i}. {q.stem}", level=2)
        if q.type == "MCQ" and q.options:
            for j, opt in enumerate(q.options):
                prefix = chr(65 + j) # A, B, C, D
                doc.add_paragraph(f"  {prefix}) {opt}")
        elif q.type == "TF":
            doc.add_paragraph("  [   ] True      [   ] False")
        elif q.type == "ShortAnswer":
            doc.add_paragraph("  __________________________________________________")
        doc.add_paragraph("")
        
    # Append separated Answer Key if role is Teacher
    if role.strip().lower() == "teacher":
        doc.add_page_break()
        doc.add_heading('Answer Key & Explanations', level=1)
        doc.add_paragraph("This section is separated for reference and evaluation purposes.")
        doc.add_paragraph("")
        
        for i, q in enumerate(quiz_data, 1):
            doc.add_heading(f"Q{i} Answer details:", level=3)
            doc.add_paragraph(f"Correct Answer: {q.correct_option}")
            doc.add_paragraph(f"Explanation: {q.answer_explanation}")
            doc.add_paragraph(f"Difficulty: {q.difficulty}")
            doc.add_paragraph("")
            
    doc.save(filepath)
    return filepath

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(80)
        self.cell(30, 10, 'Generated Quiz', 0, 0, 'C')
        self.ln(20)

def export_pdf(quiz_data: List[QuizQuestion], filepath: str, role: str = "Teacher"):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    
    # Write all questions first
    for i, q in enumerate(quiz_data, 1):
        stem = str(q.stem).encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font("Arial", "B", size=11)
        pdf.multi_cell(0, 10, text=f"Q{i}. {stem}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Arial", size=11)
        
        if q.type == "MCQ" and q.options:
            for j, opt in enumerate(q.options):
                prefix = chr(65 + j)
                opt_txt = str(opt).encode('latin-1', 'replace').decode('latin-1')
                pdf.cell(0, 8, text=f"   {prefix}) {opt_txt}", new_x="LMARGIN", new_y="NEXT")
        elif q.type == "TF":
            pdf.cell(0, 8, text="   True / False", new_x="LMARGIN", new_y="NEXT")
        elif q.type == "ShortAnswer":
            pdf.cell(0, 8, text="   Answer: _________________________________________", new_x="LMARGIN", new_y="NEXT")
            
        pdf.ln(4)
        
    # Append separated Answer Key if role is Teacher
    if role.strip().lower() == "teacher":
        pdf.add_page()
        pdf.set_font("Arial", "B", size=14)
        pdf.cell(0, 10, text="Answer Key & Explanations", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Arial", size=11)
        pdf.ln(5)
        
        for i, q in enumerate(quiz_data, 1):
            pdf.set_font("Arial", "B", size=11)
            pdf.cell(0, 8, text=f"Q{i} Answer Details:", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Arial", size=11)
            
            ans = str(q.correct_option).encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 8, text=f"   Correct Answer: {ans}", new_x="LMARGIN", new_y="NEXT")
            
            exp = str(q.answer_explanation).encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 8, text=f"   Explanation: {exp}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)

    pdf.output(filepath)
    return filepath
